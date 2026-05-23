import os
import random
from faker import Faker
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document

fake = Faker()

SKILLS_POOL = [
    "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "Go", "Rust",
    "FastAPI", "Django", "Flask", "Spring Boot", "Node.js", "React", "Vue.js",
    "Angular", "REST APIs", "GraphQL", "PostgreSQL", "MySQL", "MongoDB",
    "Redis", "Docker", "Kubernetes", "AWS", "Azure", "GCP", "Terraform",
    "Git", "Linux", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch",
    "Pandas", "NumPy", "scikit-learn", "NLP", "LangChain", "OpenAI API",
    "Kafka", "RabbitMQ", "Elasticsearch", "CI/CD", "Jenkins", "GitHub Actions",
]

JOB_TITLES = [
    "Software Engineer", "Backend Developer", "Frontend Developer",
    "Full Stack Developer", "Data Scientist", "ML Engineer",
    "DevOps Engineer", "Cloud Architect", "Python Developer",
    "Java Developer", "Data Engineer", "AI Engineer",
]

COMPANIES = [
    "Infosys", "TCS", "Wipro", "HCL Technologies", "Tech Mahindra",
    "Accenture", "Cognizant", "Capgemini", "IBM India", "Oracle India",
    "Amazon India", "Google India", "Microsoft India", "Flipkart", "Zoho",
]

DEGREES = [
    "B.Tech in Computer Science", "B.Tech in Information Technology",
    "M.Tech in Computer Science", "MCA", "B.Sc in Computer Science",
    "B.Tech in Electronics", "M.S. in Data Science","B.E in Computer Science",
]

COLLEGES = [
    "IIT Bombay", "IIT Delhi", "NIT Trichy", "VIT Vellore",
    "Anna University", "BITS Pilani", "SRM University",
    "Manipal Institute of Technology", "PSG College of Technology",
]

LOCATIONS = [
    "Bangalore", "Hyderabad", "Chennai", "Pune", "Mumbai",
    "Delhi", "Kolkata", "Noida", "Ahmedabad", "Coimbatore",
]

def safe_sample(pool, k):
    return random.sample(pool, min(k, len(pool)))

def generate_resume_data(resume_id):
    exp_years  = random.randint(0, 12)
    num_skills = random.randint(4, 10)
    num_jobs   = max(1, min(exp_years // 2, 4))
    jobs = []
    for _ in range(num_jobs):
        jobs.append({
            "title":   random.choice(JOB_TITLES),
            "company": random.choice(COMPANIES),
            "years":   f"{random.randint(2015, 2022)} - {random.randint(2022, 2024)}",
        })
    return {
        "id":          f"RES-{resume_id:05d}",
        "name":        fake.name(),
        "email":       fake.email(),
        "phone":       fake.phone_number(),
        "location":    random.choice(LOCATIONS),
        "title":       random.choice(JOB_TITLES),
        "experience":  exp_years,
        "skills":      safe_sample(SKILLS_POOL, num_skills),
        "jobs":        jobs,
        "degree":      random.choice(DEGREES),
        "college":     random.choice(COLLEGES),
        "grad_year":   random.randint(2008, 2022),
        "sandbox_url": f"https://sandbox.company.com/profiles/RES-{resume_id:05d}",
    }

def write_pdf(data, path):
    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    y = h - 50
    def line(text, size=11, bold=False, gap=18):
        nonlocal y
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        # truncate long lines to prevent overflow
        if len(text) > 100:
            text = text[:97] + "..."
        c.drawString(50, y, text)
        y -= gap

    def divider():
        nonlocal y
        c.setLineWidth(0.5)
        c.line(50, y, w - 50, y)
        y -= 10
    line(data["name"], size=16, bold=True, gap=22)
    line(data["title"], size=12, gap=16)
    line(f"{data['email']}  |  {data['phone']}  |  {data['location']}", size=9, gap=20)
    divider()
    line("PROFESSIONAL SUMMARY", size=11, bold=True, gap=16)
    summary = (f"{data['name']} is a {data['title']} with {data['experience']} years of experience "
               f"based in {data['location']}, skilled in {', '.join(data['skills'][:4])}.")
    words = summary.split()
    row = ""
    for word in words:
        if len(row) + len(word) < 88:
            row += word + " "
        else:
            line(row.strip(), size=9, gap=14)
            row = word + " "
    if row:
        line(row.strip(), size=9, gap=18)
    divider()
    line("SKILLS", size=11, bold=True, gap=16)
    line(", ".join(data["skills"]), size=9, gap=18)
    divider()
    line("EXPERIENCE", size=11, bold=True, gap=16)
    for job in data["jobs"]:
        line(f"{job['title']}  -  {job['company']}  ({job['years']})", size=9, bold=True, gap=14)
        line("  Delivered scalable backend systems, REST APIs, and cross-functional solutions.", size=9, gap=16)
    divider()

    line("EDUCATION", size=11, bold=True, gap=16)
    line(f"{data['degree']}  |  {data['college']}  |  {data['grad_year']}", size=9, gap=16)

    c.save()

def write_docx(data, path):
    doc = Document()
    doc.add_heading(data["name"], 0)
    doc.add_paragraph(data["title"])
    doc.add_paragraph(f"{data['email']}  |  {data['phone']}  |  {data['location']}")
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        f"{data['name']} is a {data['title']} with {data['experience']} years of experience "
        f"based in {data['location']}, proficient in {', '.join(data['skills'][:4])} and more."
    )
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(data["skills"]))
    doc.add_heading("Experience", level=1)
    for job in data["jobs"]:
        doc.add_paragraph(f"{job['title']} - {job['company']} ({job['years']})", style="List Bullet")
        doc.add_paragraph("Delivered backend systems, REST APIs, and scalable cloud solutions.")
    doc.add_heading("Education", level=1)
    doc.add_paragraph(f"{data['degree']}  |  {data['college']}  |  {data['grad_year']}")
    doc.save(path)

def main():
    output_dir = "resumes"
    os.makedirs(output_dir, exist_ok=True)
    total = 20  # change to however many you need
    pdf_count = docx_count = 0
    for i in range(1, total + 1):
        data = generate_resume_data(i)
        if i % 2 == 0:
            path = os.path.join(output_dir, f"{data['id']}.pdf")
            write_pdf(data, path)
            pdf_count += 1
        else:
            path = os.path.join(output_dir, f"{data['id']}.docx")
            write_docx(data, path)
            docx_count += 1
        print(f"  Created: {os.path.basename(path)}")

    print(f"\nDone — {total} resumes ({pdf_count} PDF, {docx_count} DOCX) in '{output_dir}/'")

if __name__ == "__main__":
    main()